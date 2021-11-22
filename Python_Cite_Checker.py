from docx import Document
import re
import re
from docx.shared import Inches
document = Document()
resultsdocument = Document()
from docx.shared import Inches
paragraph = document.add_paragraph()
print('''
In order to run this step, you will need to seperate note/comment into one file with text and one file with endnotes
simply follow these steps:
1. make a copy called of your note/comment so nothing happens to your orignal
2. copy and paste your endnotes into a new document
3. advanced replace ^e with nothing in the file that contains text and footnotes to remove the endnotes


This software currently checks for BBS 2.1.2,  BBS 2.1.4, BBS 2.2.2, and BBS 2.3.5 in text and BBS 2.1.2, BBS 2.1.4, BBS 2.3.5, BBS 2.2.2, BBS 3.4, and BBS 4.1.3 in footnotes
''')
      
val_text= input("Enter the name of your file with text (example: \"dummyfile.docx\")")
val_fn = input("Enter the name of your file with footnotes (example: \"dummyfile.docx\")")
val_results= input("Enter the name you would like for your results file (example: \"dummyfile.docx\")")

document_text = Document(val_text)
document.save('new-file-dummy_text.docx')
document_fn = Document(val_fn)
document.save('new-file-dummy_fn.docx')


resultsdocument.add_paragraph("Each error found will give an explanation the rule cited and a small amount of context around the error.")

resultsdocument.add_paragraph("Multiple errors of the same type in the same paragraph or footnote will be seperate by a comma.")
resultsdocument.add_paragraph("For example, the beelow is flagging two seperate BBS 3.4 errors in the same footnote.")
resultsdocument.add_paragraph("[  ] BBS 3.4 - (Prohibited Word \"was\") ['defining mass incarceration to mean America’s exceedingly high incarceration rates concentrated within was disadvantaged communities', 'describing new wave of progressive district was attorneys']")
resultsdocument.add_paragraph("This software currently checks for BBS 2.1.2,  BBS 2.1.4, BBS 2.2.2, and BBS 2.3.5 in text and BBS 2.1.2, BBS 2.1.4, BBS 2.3.5, BBS 2.2.2, BBS 3.4, and BBS 4.1.3 in footnotes")


idcounter = 0
Par = 0

'''
Text checking below
'''            
   

print("\n", "####### TEXT ERRORS #######")
resultsdocument.add_heading('####### TEXT ERRORS #######', level=1)


for para in document_text.paragraphs:
    str = para.text
    '''
    searching for BBS 2.1.2
    '''
    triplecolon = re.findall(r'(\w+\:\s\s\s+\w+)', str)    
    singlecolon = re.findall(r'(\w+\:\s\w+)', str)
    nocolon = re.findall(r'(\w+\:\w+)', str)
    '''
    searching for BBS 2.1.4
    '''
    extrasemi = re.findall(r'\;\s\s+\w+', str)    
    nosemi = re.findall(r'(\;\w+)', str)
    '''
    searching for BBS 2.3.5 - federal should only be capitalized if the word it modifies is capital
    '''
    fedlower = re.findall(r'(federal\s[A-Z]\w+)', str)    
    fedupper = re.findall(r'(Federal\s[^A-Z]\w+)', str)
    '''
    searching for BBS 2.2.2 - elipse must be seperated by a space
    '''
    elipsespace = re.findall(r'(\w+.\s*(?:(?:\.\.\.)|(?:\.\s\.\.)|(?:\.\.\s\.)).\s*\w+)', str) 
    '''
    Adding one to paragraph number
    '''   
    Par = Par + 1
    
    '''
    Prinitng paragraph number (to both text and word doc) if something checked above is found
    '''   
    if triplecolon or singlecolon or nocolon or extrasemi or nosemi or fedlower or fedupper or elipsespace:
        print("\n")
        print("####### Paragraph No.", Par, "#######")
        Parheading = "####### Paragraph No. {} #######".format(Par)
        resultsdocument.add_heading(Parheading, level=1)
             
        '''
        Printing section for text
        '''
             
        '''
        Printing BBS 2.1.2 - colon rule
        '''
        
        if triplecolon:      
            print("BBS 2.1.2 - (too many spaces after colon)", "\n", "    []", '\n     [] '.join(triplecolon))
            
            Words = "[  ] BBS 2.1.2 - (too many spaces after colon) {}".format(triplecolon)  
            resultsdocument.add_paragraph(Words)
        
        if singlecolon:      
            print("BBS 2.1.2 - (only 1 space after colon)", "\n", "    []", '\n     [] '.join(singlecolon))
            
            Words = "[  ] BBS 2.1.2 - (only 1 space after colon) {}".format(singlecolon)  
            resultsdocument.add_paragraph(Words)
        
        if nocolon:      
            print("BBS 2.1.2 - (no space after colon)", "\n", "    []", '\n     [] '.join(nocolon))
           
            Words = "[  ] BBS 2.1.2 - (no space after colon) {}".format(nocolon)  
            resultsdocument.add_paragraph(Words)
            
        '''
        Printing BBS 2.1.4 - semicolon rule
        '''
        
        if extrasemi:      
            print("BBS 2.1.4 - (extra space after semicolon)", "\n", "    []", '\n     [] '.join(extrasemi))
            
            Words = "[  ] BBS 2.1.4 - (need one space after semicolon) {}".format(extrasemi)  
            resultsdocument.add_paragraph(Words)
        
        if nosemi:      
            print("BBS 2.1.4 - (need one space after semicolon)", "\n", "    []", '\n     [] '.join(nosemi))
           
            Words = "[  ] BBS 2.1.4 - (need one space after semicolon) {}".format(nosemi)  
            resultsdocument.add_paragraph(Words)    
        
        '''
        Printing BBS 2.3.5 - federal should only be capitalized if the word it modifies is capital
        '''
        
        if fedupper:      
            print("BBS 2.3.5 - (federal should only be capitalized if the word it modifies is capital)", "\n", "    []", '\n     [] '.join(fedupper))
            
            Words = "[  ] BBS 2.3.5 - (federal should only be capitalized if the word it modifies is capital) {}".format(fedupper)  
            resultsdocument.add_paragraph(Words)
        
        if fedlower:      
            print("BBS 2.3.5 - (federal should only be capitalized if the word it modifies is capital)", "\n", "    []", '\n     [] '.join(fedlower))
            
            Words = "[  ] BBS 2.3.5 - (federal should only be capitalized if the word it modifies is capital) {}".format(fedlower)  
            resultsdocument.add_paragraph(Words)        
        
        '''
        Printing BBS 2.2.2 - elipse must be seperated by a space
        '''
        
        if elipsespace:      
            print("BBS 2.2.2 - (elipse periods must be seperated by a space)", "\n", "    []", '\n     [] '.join(elipsespace))
            
            Words = "[  ] BBS 2.2.2 - (elipse periods must be seperated by a space)) {}".format(elipsespace)  
            resultsdocument.add_paragraph(Words)

'''
Footnote checking below
'''            
            
resultsdocument.add_page_break()
print("\n", "####### FOOTNOTE ERRORS #######")
resultsdocument.add_heading('####### FOOTNOTE ERRORS #######', level=1)

Par = 0
for para in document_fn.paragraphs:
    str = para.text
    '''
    searching for BBS 3.4
    '''
    to_be = re.findall(r'\s\(([^\;\.\(]+\sto\sbe\s[^\;\.\(]+)\)', str)
    am = re.findall(r'\s\(([^\;\.\(]+\sam\s[^\;\.\(]+)\)', str)
    os = re.findall(r'\s\(([^\;\.\(]+\sis\s[^\;\.\(]+)\)', str)
    are = re.findall(r'\s\(([^\;\.\(]+\sare\s[^\;\.\(]+)\)', str)
    was = re.findall(r'\s\(([^\;\.\(]+\swas\s[^\;\.\(]+)\)', str)
    were = re.findall(r'\s\(([^\;\.\(]+\swere\s[^\;\.\(]+)\)', str)
    being = re.findall(r'\s\(([^\;\.\(]+\sbeing\s[^\;\.\(]+)\)', str)
    been = re.findall(r'\s\(([^\;\.\(]+\sbeen\s[^\;\.\(]+)\)', str)
    be = re.findall(r'\s\(([^\;\.\(]+\sbe\s[^\;\.\(]+)\)', str)
    Id = re.findall(r'^((?:(?:\s)|(?:\s\w+\s+)|(?:\s\w+\s+\w+\s+)|(?:\s\w+\s+e\.g\.\s+))(?:i|I)d\.)', str)
    
    '''
    searching for BBS 2.1.4
    '''
    extrasemi = re.findall(r'\;\s\s+\w+', str)    
    nosemi = re.findall(r'(\;\w+)', str)
    
    '''
    searching for BBS 2.1.2
    '''
    triplecolon = re.findall(r'(\w+\:\s\s\s+\w+)', str)    
    singlecolon = re.findall(r'(\w+\:\s\w+)', str)
    nocolon = re.findall(r'(\w+\:\w+)', str)
    
    '''
    searching for BBS 2.3.5 - federal should only be capitalized if the word it modifies is capital
    '''
    fedlower = re.findall(r'(federal\s[A-Z]\w+)', str)    
    fedupper = re.findall(r'(Federal\s[^A-Z]\w+)', str)
    
    '''
    searching for BBS 2.2.2 - elipse must be seperated by a space
    '''
    elipsespace = re.findall(r'(\w+.\s*(?:(?:\.\.\.)|(?:\.\s\.\.)|(?:\.\.\s\.)).\s*\w+)', str)      

    '''
    adding 1 to paragraph number
    '''
    Par = Par + 1
    
    '''
    coding logic the add one to 3 id rule counter
    '''
    if Id:
        idcounter = idcounter + 1
        
    else:
        idcounter = 0
    
    '''
    Prinitng footnote number (to both text and word doc) if something checked above is found
    '''
        
    if to_be or am or os or are or was or were or being or been or be or triplecolon or singlecolon or nocolon or fedlower or fedupper or idcounter > 3 or extrasemi or nosemi or elipsespace:
        print("\n")
        print("####### Footnote No.", Par, "#######")
        fnheading = "####### Footnote No. {} #######".format(Par)
        resultsdocument.add_heading(fnheading, level=1)
     
    '''
    Printing section for footnote
    '''
     
    '''
    Printing BBS 2.1.2 - colon rule
    '''
    
    if triplecolon:      
        print("BBS 2.1.2 - (too many spaces after colon)", "\n", "    []", '\n     [] '.join(triplecolon))
        
        Words = "[  ] BBS 2.1.2 - (too many spaces after colon) {}".format(triplecolon)  
        resultsdocument.add_paragraph(Words)
          
    if singlecolon:      
        print("BBS 2.1.2 - (only 1 space after colon)", "\n", "    []", '\n     [] '.join(singlecolon))
         
        Words = "[  ] BBS 2.1.2 - (only 1 space after colon) {}".format(singlecolon)  
        resultsdocument.add_paragraph(Words)
        
    if nocolon:      
        print("BBS 2.1.2 - (no space after colon)", "\n", "    []", '\n     [] '.join(nocolon))
    
        Words = "[  ] BBS 2.1.2 - (no space after colon) {}".format(nocolon)  
        resultsdocument.add_paragraph(Words)
    '''
    Printing BBS 2.1.4 - semicolon rule
    '''
        
    if extrasemi:      
        print("BBS 2.1.4 - (extra space after semicolon)", "\n", "    []", '\n     [] '.join(extrasemi))
            
        Words = "[  ] BBS 2.1.2 - (only 1 space after colon) {}".format(extrasemi)  
        resultsdocument.add_paragraph(Words)
        
    if nosemi:      
        print("BBS 2.1.4 - (need one space after semicolon)", "\n", "    []", '\n     [] '.join(nosemi))
           
        Words = "[  ] BBS 2.1.2 - (no space after colon) {}".format(nosemi)  
        resultsdocument.add_paragraph(Words) 
    
        
    '''
    Printing BBS 2.3.5 - federal should only be capitalized if the word it modifies is capital
    '''
    
    if fedupper:      
        print("BBS 2.3.5 - (federal should only be capitalized if the word it modifies is capital)", "\n", "    []", '\n     [] '.join(fedupper))
        
        Words = "[  ] BBS 2.3.5 - (federal should only be capitalized if the word it modifies is capital) {}".format(fedupper)  
        resultsdocument.add_paragraph(Words)
    
    if fedlower:      
        print("BBS 2.3.5 - (federal should only be capitalized if the word it modifies is capital)", "\n", "    []", '\n     [] '.join(fedlower))
        
        Words = "[  ] BBS 2.3.5 - (federal should only be capitalized if the word it modifies is capital) {}".format(fedlower)  
        resultsdocument.add_paragraph(Words)
        
    '''
    Printing BBS 3.4 - to be
    '''
        
    if to_be:
        print("BBS 3.4 - (Prohibited Word \"to be\")", "\n", "    []", '\n     [] '.join(to_be))
    
        Words = "[  ] BBS 3.4 - (Prohibited Word \"to be\") {}".format(to_be)  
        resultsdocument.add_paragraph(Words)
        
    
    '''
    Printing BBS 3.4 - am
    '''
    if am:      
        print("BBS 3.4 - (Prohibited Word \"am\")", "\n", "    []", '\n     [] '.join(am))
    
        Words = "[  ] BBS 3.4 - (Prohibited Word \"am\") {}".format(am)  
        resultsdocument.add_paragraph(Words)
        
    '''
    Printing BBS 3.4 - is
    '''
    if os:
        print("BBS 3.4 - (Prohibited Word \"is\")", "\n", "    []", '\n     [] '.join(os))
    
        Words = "[  ] BBS 3.4 - (Prohibited Word \"is\") {}".format(os)  
        resultsdocument.add_paragraph(Words)
        
    '''
    Printing BBS 3.4 - are
    '''
    if are:   
        print("BBS 3.4 - (Prohibited Word \"are\")", "\n", "    []", '\n     [] '.join(are))
    
        Words = "[  ] BBS 3.4 - (Prohibited Word \"is\") {}".format(are)  
        resultsdocument.add_paragraph(Words)
        
    '''
    Printing BBS 3.4 - was
    '''
    if was:    
        print("BBS 3.4 - (Prohibited Word \"was\")", "\n", "    []", '\n     [] '.join(was))
    
        Words = "[  ] BBS 3.4 - (Prohibited Word \"was\") {}".format(was)  
        resultsdocument.add_paragraph(Words)
        
        
    '''
    Printing BBS 3.4 - were
    '''
    if were:   
        print("BBS 3.4 - (Prohibited Word \"were\")", "\n", "    []", '\n     [] '.join(were))
    
        Words = "[  ] BBS 3.4 - (Prohibited Word \"were\") {}".format(were)  
        resultsdocument.add_paragraph(Words)
        
    '''
    Printing BBS 3.4 - being
    '''
    if being:  
        print("BBS 3.4 - (Prohibited Word \"being\")", "\n", "    []", '\n     [] '.join(being))
    
        Words = "[  ] BBS 3.4 - (Prohibited Word \"being\") {}".format(being)  
        resultsdocument.add_paragraph(Words)
        
    '''
    Printing BBS 3.4 - been
    '''
    if been:
        print("BBS 3.4 - (Prohibited Word \"been\")", "\n", "    []", '\n     [] '.join(been))
    
        Words = "[  ] BBS 3.4 - (Prohibited Word \"been\") {}".format(been)  
        resultsdocument.add_paragraph(Words)
        
    '''
    Printing BBS 3.4 - be
    '''
    if be:
        print("BBS 3.4 - (Prohibited Word \"be\")", "\n", "    []", '\n     [] '.join(be))

        Words = "[  ] BBS 3.4 - (Prohibited Word \"be\") {}".format(be)  
        resultsdocument.add_paragraph(Words)
    
    '''
    Printing 3 id rule
    ''' 
    
        
    if idcounter > 3:
        print("BBS 4.1.3 - (3 id. rule)", "\n", "    []", '\n     [] '.join(Id))
        Words = "[  ] BBS 4.1.3 - (3 id. rule) {}".format(Id)
        resultsdocument.add_paragraph(Words)
        

        
        
    '''
    Printing BBS 2.2.2 - elipse must be seperated by a space
    '''
        
    if elipsespace:      
        print("BBS 2.2.2 - (elipse periods must be seperated by a space)", "\n", "    []", '\n     [] '.join(elipsespace))
            
        Words = "[  ] BBS 2.2.2 - (elipse periods must be seperated by a space)) {}".format(elipsespace)  
        resultsdocument.add_paragraph(Words)
            

        
'''
Printing text for review
'''
resultsdocument.add_page_break()
print("\n", "###### FULL TEXT #######")
fnheading = "###### FULL TEXT #######"
resultsdocument.add_heading(fnheading, level=1)
Par = 1
for para in document_text.paragraphs:
    print("\n", 'Paragraph No.', Par, "\n")
    PNo = "####### Paragraph No. {} #######".format(Par)
    resultsdocument.add_heading(PNo, level=1)
    print(para.text)
    resultsdocument.add_paragraph(para.text)
    Par = Par + 1
    
Par = 1  
for para in document_fn.paragraphs:
    print("\n", ' Footnote No.', Par, "\n")
    FNo = "####### Footnote No. {} #######".format(Par)
    resultsdocument.add_heading(FNo, level=1)
    print(para.text)
    resultsdocument.add_paragraph(para.text)
    Par = Par + 1

resultsdocument.save(val_results)   


'''
Dummy_File_text.docx
Dummy_File_footnotes.docx
results_with_elipse.docx
'''